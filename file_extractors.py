"""
Extract text from non-PDF file types (images, DOC/DOCX, XLS/XLSX)
and handle archive extraction (ZIP, TAR, GZ).

Every extractor returns (pages, sections) in the same format as
extractor.extract_pdf() so the rest of the ingestion pipeline
(ingest_document, metadata, embeddings) works unchanged.
"""

import gzip
import shutil
import subprocess
import tarfile
import tempfile
import zipfile
from pathlib import Path

IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp', '.gif'}
DOC_EXTENSIONS = {'.doc'}
DOCX_EXTENSIONS = {'.docx'}
XLS_EXTENSIONS = {'.xls'}
XLSX_EXTENSIONS = {'.xlsx'}
PDF_EXTENSIONS = {'.pdf'}
INGESTABLE_EXTENSIONS = (
    PDF_EXTENSIONS | IMAGE_EXTENSIONS | DOC_EXTENSIONS | DOCX_EXTENSIONS | XLS_EXTENSIONS | XLSX_EXTENSIONS
)
ARCHIVE_EXTENSIONS = {'.zip', '.tar', '.tgz'}  # .tar.gz handled separately


# ---------------------------------------------------------------------------
# Archive handling
# ---------------------------------------------------------------------------

def is_archive(filename: str) -> bool:
    lower = filename.lower()
    if lower.endswith('.tar.gz'):
        return True
    return Path(lower).suffix in ARCHIVE_EXTENSIONS | {'.gz'}


def _strip_common_root(entries: list[tuple[str, bytes]]) -> list[tuple[str, bytes]]:
    """If every entry lives under a single top-level directory, strip it.

    This handles the common case where a user zips a folder and every file
    inside has the folder name as prefix (e.g. ``MyProject/specs/file.pdf``).
    """
    if not entries:
        return entries
    rel_paths = [Path(p) for p, _ in entries]
    # Only strip when *all* files are inside at least one subdirectory
    if not all(len(p.parts) > 1 for p in rel_paths):
        return entries
    roots = {p.parts[0] for p in rel_paths}
    if len(roots) != 1:
        return entries
    return [(str(Path(*Path(p).parts[1:])), d) for p, d in entries]


def _safe_archive_entry(rel_path: str) -> bool:
    """Return True if the archive entry path is safe to extract."""
    p = Path(rel_path)
    if not p.name or '..' in p.parts:
        return False
    # Skip hidden / system entries (macOS __MACOSX, .DS_Store, etc.)
    for part in p.parts:
        if part.startswith('.') or part.startswith('_'):
            return False
    if p.suffix.lower() not in INGESTABLE_EXTENSIONS:
        return False
    return True


def extract_archive(archive_path: str, dest_dir: Path) -> list[str]:
    """Extract supported files from an archive into *dest_dir*.

    Directory structure inside the archive is preserved.  If every file
    shares a single common root directory it is stripped to avoid
    double-nesting (e.g. uploading ``Project.zip`` containing
    ``Project/specs/file.pdf`` into folder "Project" won't create
    ``Project/Project/specs/file.pdf``).

    Returns list of extracted relative paths (may include subdirectories).
    """
    lower = str(archive_path).lower()

    # --- collect raw (internal_path, data) pairs -------------------------
    entries: list[tuple[str, bytes]] = []

    if lower.endswith('.zip'):
        with zipfile.ZipFile(archive_path) as zf:
            for info in zf.infolist():
                if info.is_dir():
                    continue
                entries.append((info.filename, zf.read(info.filename)))

    elif lower.endswith('.tar') or lower.endswith('.tar.gz') or lower.endswith('.tgz'):
        mode = 'r:gz' if (lower.endswith('.gz') or lower.endswith('.tgz')) else 'r'
        with tarfile.open(archive_path, mode) as tf:
            for member in tf.getmembers():
                if not member.isfile():
                    continue
                f = tf.extractfile(member)
                if f:
                    entries.append((member.name, f.read()))

    elif lower.endswith('.gz'):
        name = Path(archive_path).stem  # remove .gz
        if Path(name).suffix.lower() in INGESTABLE_EXTENSIONS:
            with gzip.open(archive_path, 'rb') as gz:
                (dest_dir / name).write_bytes(gz.read())
            return [name]
        return []

    if not entries:
        return []

    entries = _strip_common_root(entries)

    # --- write files, preserving directory structure ---------------------
    extracted: list[str] = []
    for rel_path, data in entries:
        if not _safe_archive_entry(rel_path):
            continue
        dest = dest_dir / rel_path
        # Guard against path traversal
        try:
            dest.resolve().relative_to(dest_dir.resolve())
        except ValueError:
            continue
        dest.parent.mkdir(parents=True, exist_ok=True)
        dest.write_bytes(data)
        extracted.append(Path(rel_path).as_posix())

    return extracted


# ---------------------------------------------------------------------------
# Image extraction (tesseract OCR)
# ---------------------------------------------------------------------------

def extract_image(image_path: str) -> tuple[list[dict], list[dict]]:
    """OCR an image file via tesseract -> (pages, sections)."""
    try:
        result = subprocess.run(
            ['tesseract', str(image_path), 'stdout', '-l', 'eng'],
            capture_output=True, text=True, timeout=120,
        )
        text = result.stdout.strip()
    except (FileNotFoundError, subprocess.TimeoutExpired):
        text = ''

    if not text:
        return [], []

    pages = [{'page_num': 1, 'content': text, 'breadcrumb': ''}]
    return pages, []


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------


def _convert_with_soffice(source_path: str, target_ext: str) -> Path:
    soffice = shutil.which('soffice') or shutil.which('libreoffice')
    if not soffice:
        raise RuntimeError(
            f"Legacy Office conversion for {Path(source_path).suffix.lower()} requires LibreOffice (soffice)."
        )

    with tempfile.TemporaryDirectory(prefix='office-convert-') as tmp_dir:
        result = subprocess.run(
            [
                soffice,
                '--headless',
                '--convert-to', target_ext.lstrip('.'),
                '--outdir', tmp_dir,
                str(source_path),
            ],
            capture_output=True,
            text=True,
            timeout=180,
        )
        if result.returncode != 0:
            detail = (result.stderr or result.stdout).strip() or 'unknown conversion error'
            raise RuntimeError(f"LibreOffice conversion failed: {detail}")

        converted = Path(tmp_dir) / (Path(source_path).stem + target_ext)
        if not converted.exists():
            raise RuntimeError(f"LibreOffice did not produce a {target_ext} file.")

        persistent_dir = Path(tempfile.mkdtemp(prefix='office-converted-'))
        persistent_path = persistent_dir / converted.name
        shutil.copy2(converted, persistent_path)
        return persistent_path


def _table_to_markdown(table) -> str:
    """Convert a python-docx Table to a markdown table."""
    rows = []
    for row in table.rows:
        cells = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ''

    ncols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < ncols:
            r.append('')

    lines = []
    lines.append('| ' + ' | '.join(rows[0]) + ' |')
    lines.append('| ' + ' | '.join(['---'] * ncols) + ' |')
    for r in rows[1:]:
        lines.append('| ' + ' | '.join(r) + ' |')
    return '\n'.join(lines)


def extract_docx(docx_path: str) -> tuple[list[dict], list[dict]]:
    """Extract text + tables from DOCX -> (pages, sections).

    Each heading-delimited section becomes a page.
    Tables are rendered as markdown in document order.
    """
    from docx import Document

    doc = Document(str(docx_path))

    # Map element identity -> Paragraph/Table for ordered iteration
    para_map = {id(p._element): p for p in doc.paragraphs}
    table_map = {id(t._element): t for t in doc.tables}

    pages: list[dict] = []
    sections: list[dict] = []
    current_parts: list[str] = []
    current_heading = ''

    def flush():
        nonlocal current_parts
        text = '\n'.join(current_parts).strip()
        if text:
            pages.append({
                'page_num': len(pages) + 1,
                'content': text,
                'breadcrumb': current_heading,
            })
        current_parts = []

    for child in doc.element.body:
        cid = id(child)

        if cid in para_map:
            para = para_map[cid]
            style = para.style.name if para.style else ''
            text = para.text.strip()

            if style.startswith('Heading'):
                flush()
                raw_level = style.replace('Heading', '').strip()
                level = int(raw_level) if raw_level.isdigit() else 1
                current_heading = text
                sections.append({
                    'heading': text, 'level': level,
                    'page_num': len(pages) + 1,
                    'breadcrumb': text,
                })
                if text:
                    current_parts.append(text)
            elif text:
                current_parts.append(text)

        elif cid in table_map:
            md = _table_to_markdown(table_map[cid])
            if md:
                current_parts.append('\n' + md)

    flush()

    # No headings? Chunk into ~3000-char pages
    if not sections and pages:
        all_text = '\n'.join(p['content'] for p in pages)
        pages = []
        for i in range(0, max(len(all_text), 1), 3000):
            chunk = all_text[i:i + 3000].strip()
            if chunk:
                pages.append({
                    'page_num': len(pages) + 1,
                    'content': chunk,
                    'breadcrumb': '',
                })

    return pages, sections


def extract_doc(doc_path: str) -> tuple[list[dict], list[dict]]:
    """Extract legacy Word .doc files."""
    converted_path = None
    try:
        converted_path = _convert_with_soffice(doc_path, '.docx')
        return extract_docx(str(converted_path))
    except RuntimeError as soffice_error:
        antiword = shutil.which('antiword')
        if antiword:
            result = subprocess.run(
                [antiword, str(doc_path)],
                capture_output=True,
                text=True,
                timeout=120,
            )
            text = result.stdout.strip()
            if text:
                return ([{
                    'page_num': 1,
                    'content': text,
                    'breadcrumb': '',
                }], [])
        raise RuntimeError(
            "Legacy .doc import requires LibreOffice/soffice or antiword to be installed on the server."
        ) from soffice_error
    finally:
        if converted_path:
            try:
                converted_path.unlink(missing_ok=True)
                converted_path.parent.rmdir()
            except OSError:
                pass


# ---------------------------------------------------------------------------
# XLS / XLSX extraction
# ---------------------------------------------------------------------------


def _rows_to_sheet_page(sheet_name: str, rows: list[list[str]], page_num: int) -> dict:
    if not rows:
        return {
            'page_num': page_num,
            'content': f'{sheet_name}\n[Empty sheet]',
            'breadcrumb': sheet_name,
        }

    ncols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < ncols:
            r.append('')

    lines = [sheet_name, '']
    lines.append('| ' + ' | '.join(rows[0]) + ' |')
    lines.append('| ' + ' | '.join(['---'] * ncols) + ' |')
    for r in rows[1:]:
        lines.append('| ' + ' | '.join(r) + ' |')

    return {
        'page_num': page_num,
        'content': '\n'.join(lines),
        'breadcrumb': sheet_name,
    }


def extract_xlsx(xlsx_path: str) -> tuple[list[dict], list[dict]]:
    """Extract sheets from XLSX -> (pages, sections)."""
    from openpyxl import load_workbook

    wb = load_workbook(str(xlsx_path), data_only=True, read_only=True)
    pages: list[dict] = []
    sections: list[dict] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        page_num = len(pages) + 1
        sections.append({
            'heading': sheet_name, 'level': 1,
            'page_num': page_num, 'breadcrumb': sheet_name,
        })

        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() if c is not None else '' for c in row]
            if any(cells):
                rows.append(cells)

        pages.append(_rows_to_sheet_page(sheet_name, rows, page_num))

    wb.close()
    return pages, sections


def extract_xls(xls_path: str) -> tuple[list[dict], list[dict]]:
    """Extract sheets from legacy XLS -> (pages, sections)."""
    import xlrd

    wb = xlrd.open_workbook(str(xls_path), on_demand=True)
    pages: list[dict] = []
    sections: list[dict] = []

    for sheet_name in wb.sheet_names():
        ws = wb.sheet_by_name(sheet_name)
        page_num = len(pages) + 1
        sections.append({
            'heading': sheet_name, 'level': 1,
            'page_num': page_num, 'breadcrumb': sheet_name,
        })

        rows = []
        for row_idx in range(ws.nrows):
            row = []
            for value in ws.row_values(row_idx):
                if value is None:
                    row.append('')
                elif isinstance(value, float) and value.is_integer():
                    row.append(str(int(value)))
                else:
                    row.append(str(value).strip())
            if any(row):
                rows.append(row)

        pages.append(_rows_to_sheet_page(sheet_name, rows, page_num))

    wb.release_resources()
    return pages, sections


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

def extract_file(file_path: str) -> tuple[list[dict], list[dict]]:
    """Route to the right extractor based on file extension.

    PDFs go through the full extractor.extract_pdf() pipeline.
    Other types use the simpler extractors above.
    """
    ext = Path(file_path).suffix.lower()

    if ext in PDF_EXTENSIONS:
        from extractor import extract_pdf
        return extract_pdf(file_path)
    if ext in IMAGE_EXTENSIONS:
        return extract_image(file_path)
    if ext in DOC_EXTENSIONS:
        return extract_doc(file_path)
    if ext in DOCX_EXTENSIONS:
        return extract_docx(file_path)
    if ext in XLS_EXTENSIONS:
        return extract_xls(file_path)
    if ext in XLSX_EXTENSIONS:
        return extract_xlsx(file_path)

    print(f"  Unsupported file type: {ext}")
    return [], []
